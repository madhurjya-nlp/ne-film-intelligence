#!/usr/bin/env python3
"""Generate article-style research docs (docx) for cinema-edu.
1. Cultural context Assam/NE India cinema & media.
2. Practical pathways abroad for Assam/NE students/filmmakers - ground realities.
Run after pip install python-docx. Outputs in research/ folder.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading_with_style(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h

# ===== DOC 1: Cultural Context =====
doc1 = Document()

# Set narrow margins
for section in doc1.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

title = doc1.add_heading('Cultural Context: Media & Cinema in Assam and Northeast India', 0)
subtitle = doc1.add_paragraph()
subtitle.add_run('Research Notes for Aspiring Filmmakers & Students from the Region').italic = True
subtitle.add_run('\nPrepared: June 2026 | For decision-making when exploring education paths (local + abroad)')

doc1.add_paragraph()

# Intro
doc1.add_heading('Why This Matters', level=1)
p = doc1.add_paragraph()
p.add_run('Before choosing universities or programs abroad, understand the ecosystem you come from. Assam and the broader Northeast (eight states) have a distinct cinematic and media identity — slow-paced, culturally rooted, often marginalized in national narratives, but rich with stories the world is increasingly seeking. Your background is not a limitation; it is a differentiator. This document summarizes history, current ground reality, challenges, and opportunities so you can decide with clarity what kind of training (local, hybrid, abroad) best serves your voice and practical constraints (distance from North Lakhimpur, ST status, budget, language).')

# History
doc1.add_heading('Historical Roots', level=1)
bullets = [
    "1935: Birth of Assamese cinema with 'Joymoti' directed by Rupkonwar Jyoti Prasad Agarwala (poet, playwright, freedom fighter, tea planter). Shot using a temporary studio 'Chitraban' at Bholaguri Tea Estate. First Assamese talkie feature.",
    "Style: Emerged as sensitive, slow-paced, literary-influenced — focused on social themes, Assamese culture, anti-colonial sentiment. Early films like 'Indramalati' (1939).",
    "Jyoti Chitraban Film Studio (Guwahati): Established 1961 by Govt of Assam (functional ~1968) in Kahilipara, named in tribute to Agarwala. Only full-fledged studio in the entire Northeast for decades. Hub for production, now also festivals and training.",
    "Dr Bhupen Hazarika Regional Government Film & Television Institute (DBHRGFTI, formerly JCFTI): Key training institute located near Guwahati. Offers 3-year diplomas in Cinematography, Editing, Sound Engineering and Acting certificate. Practical, affordable, government-run — the realistic entry point for most from Assam and NE.",
]
for b in bullets:
    doc1.add_paragraph(b, style='List Bullet')

# Current Scene
doc1.add_heading('Current Scene (2025-2026)', level=1)
doc1.add_paragraph("Independent cinema is the vibrant core. Commercial 'Jollywood' exists but is limited in scale compared to South Indian industries. Growth areas:")

curr = [
    "Festivals & Visibility: Northeast India Film Festival (NEIFF, Manipur), Northeast India International Film Festival (NIIFF), Shillong International Film Festival, special NE sections at MIFF and national events. Films screening at New York Indian Film Festival, international shorts circuits.",
    "New Voices: Young directors (e.g. Snigdha P. Roy 'Aakuti', Maharshi Tuhin Kashyap, others) bringing fresh narratives. Language films beyond Assamese — Tai Phake, Bodo, Mising experiments for authentic representation.",
    "OTT & Digital: Increasing interest from platforms. Short-form + documentaries travel better. Reels and regional content creators building audiences.",
    "Studio Evolution: Jyoti Chitraban remains symbolic and practical center — hosts Chalachitram National Film Festival, NE Documentary Fest. Transitioning analog heritage to digital production support.",
]
for c in curr:
    doc1.add_paragraph(c, style='List Bullet')

# Challenges Ground Reality
doc1.add_heading('Ground Reality: Challenges for Filmmakers from Assam/NE', level=1)
chal = [
    "Infrastructure: Single major studio for 8 states historically. Equipment, post-production often requires travel to Guwahati/Kolkata/Mumbai. Power/internet reliability issues in smaller towns like Lakhimpur.",
    "Funding & Distribution: Limited local producers, weak theatrical in region. National distributors favor Hindi/regional big languages. Subtitles + cultural translation barriers for wider Indian/international reach.",
    "Representation: Historical mainstream Assamese cinema sometimes stereotyped or marginalized indigenous tribal experiences (Bodo, Mising, Dimasa etc.). Independent filmmakers increasingly reclaiming narrative sovereignty.",
    "Scale & Mobility: From remote Assam, attending workshops/fests in Guwahati adds cost/time. Going to metros or abroad requires significant planning (family, finances, ST certificate logistics).",
    "Language: Assamese (and 100+ NE languages) rich but niche. English proficiency + subtitles essential for global/education paths.",
]
for ch in chal:
    doc1.add_paragraph(ch, style='List Bullet')

# Opportunities
doc1.add_heading('Opportunities & Strategic Fit', level=1)
opp = [
    "Global hunger for authentic, diverse, indigenous, ecological, and postcolonial stories. NE India (biodiversity, Brahmaputra, tribal oral traditions, Bihu, post-insurgency transitions, tea gardens, unique geopolitics) is underrepresented — your advantage.",
    "ST status: Real leverage for National Overseas Scholarship (NOS), state schemes, and some international funds that prioritize marginalized voices (e.g. Hubert Bals prioritizes artistic projects from Global South that struggle locally).",
    "Local-to-Global Pipeline: Start with DBHRGFTI or self-shot work + NE fests → build showreel → use NOS/Erasmus for advanced craft/networks → return or co-produce with strengthened profile.",
    "Hybrid Power: Many successful regional filmmakers combine deep local roots with selective international exposure (workshops, short programs, or full degrees).",
]
for o in opp:
    doc1.add_paragraph(o, style='List Bullet')

# Decision Guidance
doc1.add_heading('How This Informs Education Choices', level=1)
doc1.add_paragraph('If your goal is to tell Assamese/NE stories authentically:')
dec = [
    "Prioritize programs with documentary/non-fiction strengths or strong indigenous/postcolonial film theory (DocNomads, some European media studies, EICTV workshops).",
    "Consider affordable/local first (DBHRGFTI + online certs + festivals) to test voice and build 3-5 finished works before expensive moves.",
    "When going abroad: Choose places with scholarship paths (NOS covers top QS) or very low tuition + living (Romania, Germany admin-fee only, Georgia, Baltic). Factor return ticket and attestation costs from Guwahati.",
    "Avoid purely commercial Hollywood-style programs if your core is slow cinema, cultural specificity, or docs — unless you want to learn craft to subvert it.",
]
for d in dec:
    doc1.add_paragraph(d, style='List Bullet')

doc1.add_paragraph()
p = doc1.add_paragraph()
p.add_run('Sources: ').bold = True
p.add_run('Wikipedia Assamese cinema, Jyoti Chitraban official, DBHRGFTI site, academic papers on representational tides & indigenous voices (2025-26), festival reports (NEIFF, NIIFF, MIFF), ground reporting from regional filmmakers, ResearchGate/Academia summaries. Always cross-verify dates and opportunities.')

doc1.add_paragraph('Last updated June 2026. For personal research use.')

doc1.save('Cultural_Context_Assam_NE_Cinema_Media.docx')
print("Created: Cultural_Context_Assam_NE_Cinema_Media.docx")

# ===== DOC 2: Pathways Abroad - Ground Realities =====
doc2 = Document()
for section in doc2.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

doc2.add_heading('Practical Pathways Abroad for Students & Filmmakers from Assam, Northeast India', 0)
p = doc2.add_paragraph()
p.add_run('Ground Reality Guide — Budget, Logistics, Scholarships, Visas, Cultural Fit | June 2026').italic = True

doc2.add_heading('Profile Context (Typical)', level=1)
ctx = doc2.add_paragraph()
ctx.add_run('25yo ST candidate from North Lakhimpur, Assam. Documentary/short filmmaker (some Doordarshan work), drone/editing skills, runs small creative agency. Budget comfortable 0-5L self; up to 10-15L with scholarship. Fluent Assamese + English. Limited foreign language. Wants formal training or credentials without losing cultural voice or going into debt.')

doc2.add_heading('1. The Biggest Lever: National Overseas Scholarship (NOS) for ST', level=1)
nos = [
    "Admin: Ministry of Tribal Affairs. Portal: overseas.tribal.gov.in",
    "What it covers (huge): Full tuition at eligible foreign university (top QS ranks preferred), monthly maintenance (~$15,400/yr US equiv), visa, medical, airfare both ways.",
    "Eligibility snapshot (verify current): ST certificate, family income <=6L, age typically <=32 for MA, min 55% in qualifying degree, unconditional offer or strong application to QS-listed uni. Film/Media covered under Fine Arts/Humanities/Social Sciences.",
    "Deadline reality (2026-27): Applications were open into late June 2026 — act immediately or monitor for next cycle. Prepare documents (ST, income, marksheets, passport, offer letters) in advance. Guwahati attestation helps.",
    "Ground note: No bond after graduation (unlike some older schemes). Highly competitive — strong portfolio + motivation letter tying your NE stories to global relevance wins.",
]
for n in nos:
    doc2.add_paragraph(n, style='List Bullet')

doc2.add_heading('2. Strong Complementary Paths (Not Mutually Exclusive)', level=1)
paths = [
    "Erasmus Mundus Joint Masters (film-specific): DocNomads (doc), KinoEyes (fiction), FilmMemory, MediaAC, ReSound. Fully funded + stipend if selected. Apply up to 3. Deadlines Jan (scholarship). Multi-country experience. English. Perfect for NE stories on world stage.",
    "Low-tuition Europe: Sapientia Romania (~€550-2k/yr), German public unis (0 tuition + small admin for many Media MA), select Poland (~€2-6k/yr English), Georgia (very cheap country + programs ~$2.5-4.5k/yr), Baltic (Estonia/Latvia/Lithuania low-mid + Erasmus links).",
    "India-first smart stack: DBHRGFTI diploma (local, cheap, practical) + Bir Tikendrajit or other affordable online MA + high-value certs (EICTV, Sundance Collab with aid, Coursera financial aid) + festival submissions. Then apply abroad with real work.",
    "Asia scholarships: Global Korea Scholarship (GKS) — full ride but Korean often needed. Taiwan MOE/ICDF. Malaysia online flexible options.",
]
for pa in paths:
    doc2.add_paragraph(pa, style='List Bullet')

doc2.add_heading('3. Ground Realities Specific to Assam/NE Going Abroad', level=1)
gr = [
    "Geography & Logistics: North Lakhimpur → Guwahati (train/bus 6-12 hrs, or flight via GAU). Most embassies/VFS in Delhi, some Kolkata. Plan 2-3 trips for attestation, biometrics, interviews. Add ₹15-40k+ for domestic travel.",
    "Airports & Flights: Fly GAU direct to Delhi/Kolkata/Bangkok/Dubai then onwards. Europe one-way often ₹50-85k. Book early or use scholar travel allowance.",
    "Documentation from NE: ST certificate (district level, update if needed), degree attestation (Guwahati HRD → MEA), police verification. Start 3-6 months early. Internet can be patchy — download forms, keep physical copies.",
    "Tests: IELTS (available Guwahati or nearby centers). Target 6.5 overall. Some programs accept MOI (medium of instruction letter) but not always.",
    "Financial Proof: NOS or offer of full scholarship removes most burden. Otherwise show bank statements, sponsor (family), education loan. German blocked account ~€11k+ for visa.",
    "Post-Arrival Practical: Part-time work allowed (Germany 20h/wk, others vary). Student networks important. Cold weather shock from Assam. Food/cultural adjustment real — seek Indian student associations.",
    "Recognition back home: Degrees from reputable unis generally fine via equivalence. NOS returnees have good credibility for jobs, teaching, grants in India + NE specific schemes.",
]
for g in gr:
    doc2.add_paragraph(g, style='List Bullet')

doc2.add_heading('4. Recommended Decision Framework (Keep Ground Reality)', level=1)
fw = [
    "Phase 0 (Now, ₹0-1L): Master free/cheap tools (DaVinci, Blender), complete 2-3 shorts/docs from local stories, get Google/Adobe certs, enroll DBHRGFTI or equivalent local if possible, submit to 2-3 NE fests.",
    "Phase 1 (Apply, low risk): NOS + 2-3 Erasmus Mundus. Also cheap backups (Sapientia, German unis, Georgia).",
    "Phase 2 (If funded 0 cost): Go for multi-country practical MA (DocNomads etc).",
    "Phase 2 alternative (4-7L self or partial): Strong Indian (SRFTI) or low-Europe + workshops.",
    "Never: Go into heavy debt for a film degree without scholarship/portfolio proof. Portfolio and completed work matter far more than the name of the school.",
]
for f in fw:
    doc2.add_paragraph(f, style='List Bullet')

doc2.add_heading('5. Key Deadlines & Next Actions (as of mid-June 2026)', level=1)
dl = [
    "NOS: Monitor overseas.tribal.gov.in — recent cycle closed ~June 30. Prepare for next immediately.",
    "Erasmus Mundus: Catalog opens ~Oct 2026 for Sept 2027 intake. Deadlines Jan 2027 scholarship. Start shortlisting now (DocNomads top priority for docs).",
    "IELTS / offers: Book test, research specific uni deadlines (many rolling or March-June).",
    "Portfolio: Continuous. Have 3 finished pieces + treatment for one feature/doc idea.",
]
for d in dl:
    doc2.add_paragraph(d, style='List Bullet')

p = doc2.add_paragraph()
p.add_run('\nDisclaimer: ').bold = True
p.add_run('Costs converted approx (1 EUR ~ ₹90, 1 USD ~ ₹83). Always verify official sites. Living costs, exchange rates, and policies change. This is synthesized research for personal navigation, not official advice.')

doc2.add_paragraph('Sources include Ministry portals, university sites, study portals (mastersportal, study.eu, studyinestonia etc), recent festival and filmmaker reports, AGENT_MEMORY context for profile.')

doc2.save('Pathways_Abroad_Assam_NE_Ground_Reality.docx')
print("Created: Pathways_Abroad_Assam_NE_Ground_Reality.docx")

print("\nBoth article-style research docs generated successfully.")